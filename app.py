import streamlit as st
import requests
from io import BytesIO
from collections import Counter
from datetime import datetime

# PDF Generation
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyleimport streamlit as st
import requests
from io import BytesIO
from collections import Counter
from datetime import datetime

# PDF Generation
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import cm
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak, Table, TableStyle
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY, TA_LEFT

# DOCX Generation
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ============================================
# KONFIGURATION
# ============================================
try:
    API_BASE_URL = st.secrets["API_BASE_URL"]
except:
    API_BASE_URL = "http://217.154.156.197:8003"

MODELS_ENDPOINT = f"{API_BASE_URL}/desInfo/models"
ANALYZE_ENDPOINT = f"{API_BASE_URL}/desInfo/generateReport"

# Kategorie Scoring
CATEGORY_POINTS = {
    'FALSCH': 5,
    'DELEGITIMIERUNG': 4,
    'VERZERRUNG': 3,
    'FRAME': 1,
    'WAHR': 0
}

# Score-Grade Mapping
SCORE_GRADES = {
    'A': (0.0, 0.4, 'wahr', 'Die Aussagen entsprechen nachprüfbaren Fakten und sind wahrheitsgetreu.'),
    'B': (0.5, 1.4, 'geframed', 'Die Aussagen enthalten wahre Kernaussagen, die durch normative Framings eingefärbt werden.'),
    'C': (1.5, 2.4, 'verzerrend', 'Die Aussagen enthalten formal richtige Kernaussagen, die durch Übertreibung, Verkürzung oder Kontextausblendung ein schiefes Bild erzeugen.'),
    'D': (2.5, 3.4, 'demokratisch dysfunktional', 'Die Aussagen enthalten substanzielle Falschbehauptungen und delegitimierende Elemente, die den demokratischen Diskurs belasten.'),
    'E': (3.5, 5.0, 'demokratisch destruktiv', 'Die Aussagen sind überwiegend falsch und delegitimierend, untergraben systematisch Vertrauen und demokratische Institutionen.')
}

CATEGORY_DESCRIPTIONS = {
    'FALSCH': 'Objektiv widerlegte Behauptungen. Es liegen belastbare Daten oder Ereignisprotokolle vor, die das Gegenteil zeigen. Keine Meinungen oder Prognosen; nur Tatsachenbehauptungen, die nachweislich nicht stimmen.',
    'DELEGITIMIERUNG': 'Abwertung oder Untergrabung von Personen, Gruppen oder Institutionen. Sprachliche Diskreditierungen, Kampfbegriffe oder systematische Diffamierungen.',
    'VERZERRUNG': 'Formal richtige Kernaussagen, die durch Übertreibung, Verkürzung oder Kontextausblendung ein schiefes Bild erzeugen.',
    'FRAME': 'Sprachliche Deutungsrahmen, die neutrale Sachverhalte emotional oder normativ aufladen, ohne sie faktisch falsch oder delegitimierend darzustellen.',
    'WAHR': 'Sachlich zutreffende Aussagen, die überprüfbar sind und keinen überzogenen Frame, keine Verzerrung oder Delegitimierung enthalten.'
}

# ============================================
# PAGE CONFIG
# ============================================
st.set_page_config(
    page_title="DESINFO Analyzer",
    page_icon="📊",
    layout="wide",
    initial_sidebar_state="expanded"
)

# ============================================
# CUSTOM CSS - Fixed Colors
# ============================================
st.markdown("""
<style>
    @import url('https://fonts.googleapis.com/css2?family=Inter:wght@400;600;700;800&display=swap');
    
    * {
        font-family: 'Inter', -apple-system, BlinkMacSystemFont, sans-serif;
    }
    
    /* Main Container */
    .main {
        background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
    }
    
    .block-container {
        max-width: 1400px;
        padding: 2rem 1rem;
    }
    
    /* Sidebar */
    [data-testid="stSidebar"] {
        background: white;
    }
    
    /* Header */
    .main-header {
        font-size: 3.5rem;
        font-weight: 800;
        color: white;
        text-align: center;
        padding: 2rem 0 0.5rem 0;
        text-shadow: 0 2px 10px rgba(0,0,0,0.2);
    }
    
    .subtitle {
        font-size: 1.5rem;
        color: #ffd93d;
        text-align: center;
        font-weight: 600;
        margin-bottom: 2rem;
    }
    
    /* White Content Box */
    .content-box {
        background: white;
        border-radius: 20px;
        padding: 2.5rem;
        margin: 1.5rem 0;
        box-shadow: 0 10px 40px rgba(0,0,0,0.3);
    }
    
    /* Section Headers */
    .section-title {
        font-size: 2rem;
        font-weight: 700;
        color: #1a1a1a;
        margin-bottom: 1.5rem;
        text-align: center;
    }
    
    /* Buttons */
    .stButton>button {
        background: linear-gradient(135deg, #ffd93d 0%, #ffb800 100%);
        color: #1a1a1a;
        border: none;
        padding: 0.75rem 2rem;
        font-size: 1.1rem;
        font-weight: 700;
        border-radius: 50px;
        transition: all 0.3s;
        width: 100%;
    }
    
    .stButton>button:hover {
        transform: translateY(-2px);
        box-shadow: 0 6px 20px rgba(255, 217, 61, 0.4);
    }
    
    /* Text Area */
    .stTextArea textarea {
        border: 2px solid #667eea;
        border-radius: 15px;
        font-size: 1rem;
        padding: 1rem;
    }
    
    /* Score Box */
    .score-box {
        background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
        border-radius: 20px;
        padding: 3rem 2rem;
        text-align: center;
        color: white;
        margin: 2rem 0;
        box-shadow: 0 10px 30px rgba(102, 126, 234, 0.4);
    }
    
    .score-value {
        font-size: 6rem;
        font-weight: 800;
        margin: 0;
    }
    
    .score-grade {
        font-size: 2.5rem;
        font-weight: 700;
        margin: 1rem 0;
        color: #ffd93d;
    }
    
    .score-label {
        font-size: 1.8rem;
        font-weight: 600;
    }
    
    /* Metric Cards */
    .metric-card {
        text-align: center;
        padding: 1.5rem 1rem;
        border-radius: 15px;
        color: white;
        margin: 0.5rem;
        box-shadow: 0 4px 12px rgba(0,0,0,0.2);
    }
    
    .metric-value {
        font-size: 2.5rem;
        font-weight: 800;
        margin: 0.3rem 0;
    }
    
    .metric-label {
        font-size: 1rem;
        font-weight: 600;
    }
    
    .metric-percent {
        font-size: 0.9rem;
        opacity: 0.9;
    }
    
    /* Statement Card */
    .statement-card {
        background: #f8f9fa;
        border-left: 5px solid;
        padding: 1.5rem;
        margin: 1.5rem 0;
        border-radius: 10px;
        box-shadow: 0 2px 8px rgba(0,0,0,0.1);
    }
    
    .statement-title {
        font-size: 1.1rem;
        font-weight: 700;
        color: #1a1a1a;
        margin-bottom: 0.8rem;
    }
    
    .statement-text {
        font-size: 1rem;
        color: #333;
        line-height: 1.6;
    }
    
    /* Category Badges */
    .category-badge {
        display: inline-block;
        padding: 0.4rem 1rem;
        border-radius: 20px;
        font-weight: 700;
        font-size: 0.85rem;
        margin: 0.3rem 0.2rem;
        color: white;
    }
    
    .badge-falsch { background: #dc3545; }
    .badge-delegitimierung { background: #fd7e14; }
    .badge-verzerrung { background: #ffc107; color: #000; }
    .badge-frame { background: #28a745; }
    .badge-wahr { background: #007bff; }
    
    /* Category Header */
    .category-header {
        font-size: 1.8rem;
        font-weight: 700;
        color: #1a1a1a;
        margin: 2rem 0 1rem 0;
        padding-bottom: 0.5rem;
        border-bottom: 3px solid #667eea;
    }
    
    .category-description {
        font-style: italic;
        color: #666;
        margin-bottom: 1.5rem;
        padding: 1rem;
        background: #f0f0f0;
        border-radius: 8px;
    }
    
    /* Hide Streamlit Elements */
    #MainMenu {visibility: hidden;}
    footer {visibility: hidden;}
    .stDeployButton {display:none;}
</style>
""", unsafe_allow_html=True)

# ============================================
# HELPER FUNCTIONS
# ============================================

@st.cache_data(ttl=300)
def fetch_models():
    try:
        response = requests.get(MODELS_ENDPOINT, timeout=10)
        if response.ok:
            data = response.json()
            return [m for m in data.get('models', []) if m.get('valid', False)]
    except Exception as e:
        st.error(f"⚠️ Fehler beim Laden der Models: {e}")
    return []

def parse_statements(text: str) -> list:
    if '|' in text:
        statements = [s.strip() for s in text.split('|') if s.strip()]
    else:
        statements = [s.strip() for s in text.split('\n') if s.strip()]
    return statements

def call_api(statements: list, model_id: int) -> dict:
    text_param = "|".join(statements)
    params = {"modelID": model_id, "text": text_param}
    
    try:
        response = requests.get(ANALYZE_ENDPOINT, params=params, timeout=120)
        response.raise_for_status()
        return {"success": True, "data": response.json()}
    except Exception as e:
        return {"success": False, "error": str(e)}

def calculate_summary(analysis_data: list) -> dict:
    category_counts = Counter(item['kategorie'] for item in analysis_data)
    total_points = sum(CATEGORY_POINTS.get(item['kategorie'], 0) for item in analysis_data)
    total_statements = len(analysis_data)
    desinfo_score = total_points / total_statements if total_statements > 0 else 0
    
    grade = 'E'
    grade_label = ''
    grade_description = ''
    
    for g, (min_score, max_score, label, desc) in SCORE_GRADES.items():
        if min_score <= desinfo_score <= max_score:
            grade = g
            grade_label = label
            grade_description = desc
            break
    
    return {
        'total_statements': total_statements,
        'category_counts': dict(category_counts),
        'total_points': total_points,
        'desinfo_score': round(desinfo_score, 1),
        'grade': grade,
        'grade_label': grade_label,
        'grade_description': grade_description
    }

def get_category_color(category: str) -> tuple:
    color_map = {
        'FALSCH': ('#dc3545', colors.HexColor('#dc3545')),
        'DELEGITIMIERUNG': ('#fd7e14', colors.HexColor('#fd7e14')),
        'VERZERRUNG': ('#ffc107', colors.HexColor('#ffc107')),
        'FRAME': ('#28a745', colors.HexColor('#28a745')),
        'WAHR': ('#007bff', colors.HexColor('#007bff'))
    }
    return color_map.get(category, ('#6c757d', colors.grey))

def generate_pdf_report(analysis_data: list, summary: dict, model_info: dict = None) -> BytesIO:
    """Generiert PDF Report - EXAKT wie Beispiel"""
    buffer = BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=A4,
        topMargin=2*cm,
        bottomMargin=2*cm,
        leftMargin=2.5*cm,
        rightMargin=2.5*cm
    )
    
    elements = []
    styles = getSampleStyleSheet()
    
    title_style = ParagraphStyle(
        'Title',
        parent=styles['Heading1'],
        fontSize=24,
        leading=28,
        textColor=colors.black,
        spaceAfter=20,
        alignment=TA_LEFT,
        fontName='Helvetica-Bold'
    )
    
    heading_style = ParagraphStyle(
        'Heading',
        parent=styles['Heading2'],
        fontSize=16,
        leading=20,
        textColor=colors.black,
        spaceAfter=10,
        spaceBefore=15,
        fontName='Helvetica-Bold'
    )
    
    body_style = ParagraphStyle(
        'Body',
        parent=styles['Normal'],
        fontSize=11,
        leading=15,
        alignment=TA_JUSTIFY,
        fontName='Helvetica'
    )
    
    italic_style = ParagraphStyle(
        'Italic',
        parent=body_style,
        fontName='Helvetica-Oblique'
    )
    
    # Titel
    elements.append(Paragraph("Vollständige Auswertung", title_style))
    elements.append(Spacer(1, 0.3*cm))
    
    # Linie
    line = Table([['']], colWidths=[16*cm])
    line.setStyle(TableStyle([('LINEABOVE', (0,0), (-1,0), 1, colors.grey)]))
    elements.append(line)
    elements.append(Spacer(1, 0.8*cm))
    
    # Zusammenfassung
    elements.append(Paragraph("Zusammenfassung", heading_style))
    total = summary['total_statements']
    counts = summary['category_counts']
    
    summary_text = f"Die vorliegende Analyse untersucht {total} Aussagen nach ihrer faktischen Richtigkeit und kommunikativen Qualität. Das Ergebnis zeigt eine ausgeprägte Tendenz zu {summary['grade_label']}en Darstellungen."
    
    elements.append(Paragraph(summary_text, body_style))
    elements.append(Spacer(1, 0.8*cm))
    elements.append(line)
    elements.append(Spacer(1, 0.8*cm))
    
    # Quantifizierung
    elements.append(Paragraph("Quantifizierung", heading_style))
    elements.append(Paragraph(f"Anzahl Aussagen gesamt: {total}", body_style))
    elements.append(Spacer(1, 0.3*cm))
    
    for cat in ['FALSCH', 'DELEGITIMIERUNG', 'VERZERRUNG', 'FRAME', 'WAHR']:
        count = counts.get(cat, 0)
        percentage = (count / total * 100) if total > 0 else 0
        hex_color, _ = get_category_color(cat)
        cat_text = f'<font color="{hex_color}">■</font> {cat}: {count} ({percentage:.0f} %)'
        elements.append(Paragraph(cat_text, body_style))
    
    elements.append(Spacer(1, 0.8*cm))
    elements.append(line)
    elements.append(Spacer(1, 0.8*cm))
    
    # Scoring
    elements.append(Paragraph("Scoring", heading_style))
    score_text = f"<b>Desinfo-Score: {summary['desinfo_score']}</b><br/>{summary['grade']}: {summary['grade_label']}<br/><br/>{summary['grade_description']}"
    elements.append(Paragraph(score_text, body_style))
    elements.append(Spacer(1, 0.8*cm))
    elements.append(line)
    elements.append(Spacer(1, 0.8*cm))
    
    # Einteilung
    elements.append(Paragraph("Einteilung", heading_style))
    for grade, (min_s, max_s, label, _) in SCORE_GRADES.items():
        elements.append(Paragraph(f"{grade}: {min_s} bis {max_s}: {label}", body_style))
    
    elements.append(PageBreak())
    
    # Vollständige Auswertung
    elements.append(Paragraph("Vollständige Auswertung des Textes", title_style))
    elements.append(Spacer(1, 0.3*cm))
    elements.append(line)
    elements.append(Spacer(1, 0.8*cm))
    
    categories_order = ['FALSCH', 'DELEGITIMIERUNG', 'VERZERRUNG', 'FRAME', 'WAHR']
    
    for category in categories_order:
        cat_items = [item for item in analysis_data if item['kategorie'] == category]
        if not cat_items:
            continue
        
        elements.append(Paragraph(f"<b>{category} ({len(cat_items)})</b>", heading_style))
        elements.append(Paragraph(CATEGORY_DESCRIPTIONS[category], italic_style))
        elements.append(Spacer(1, 0.5*cm))
        
        for idx, item in enumerate(cat_items, 1):
            hex_color, _ = get_category_color(category)
            symbol = '✓' if category == 'WAHR' else '■'
            
            stmt_text = f'{idx}. <font color="{hex_color}">{symbol}</font> <b>"{item["aussage"]}"</b>'
            elements.append(Paragraph(stmt_text, body_style))
            elements.append(Spacer(1, 0.2*cm))
            
            beg_text = f"– {item['begründung']}"
            elements.append(Paragraph(beg_text, body_style))
            elements.append(Spacer(1, 0.5*cm))
        
        elements.append(Spacer(1, 0.5*cm))
    
    doc.build(elements)
    buffer.seek(0)
    return buffer

def generate_docx_report(analysis_data: list, summary: dict, model_info: dict = None) -> BytesIO:
    """Generiert DOCX Report"""
    doc = Document()
    
    doc.add_heading('Vollständige Auswertung', 0)
    doc.add_heading('Zusammenfassung', 1)
    
    total = summary['total_statements']
    summary_text = f"Die vorliegende Analyse untersucht {total} Aussagen nach ihrer faktischen Richtigkeit und kommunikativen Qualität."
    doc.add_paragraph(summary_text)
    
    doc.add_heading('Quantifizierung', 1)
    doc.add_paragraph(f"Anzahl Aussagen gesamt: {total}")
    
    counts = summary['category_counts']
    for cat in ['FALSCH', 'DELEGITIMIERUNG', 'VERZERRUNG', 'FRAME', 'WAHR']:
        count = counts.get(cat, 0)
        percentage = (count / total * 100) if total > 0 else 0
        doc.add_paragraph(f"■ {cat}: {count} ({percentage:.0f} %)")
    
    doc.add_heading('Scoring', 1)
    doc.add_paragraph(f"Desinfo-Score: {summary['desinfo_score']}")
    doc.add_paragraph(f"{summary['grade']}: {summary['grade_label']}")
    doc.add_paragraph(summary['grade_description'])
    
    doc.add_heading('Einteilung', 1)
    for grade, (min_s, max_s, label, _) in SCORE_GRADES.items():
        doc.add_paragraph(f"{grade}: {min_s} bis {max_s}: {label}")
    
    doc.add_page_break()
    doc.add_heading('Vollständige Auswertung des Textes', 1)
    
    categories_order = ['FALSCH', 'DELEGITIMIERUNG', 'VERZERRUNG', 'FRAME', 'WAHR']
    
    for category in categories_order:
        cat_items = [item for item in analysis_data if item['kategorie'] == category]
        if not cat_items:
            continue
        
        doc.add_heading(f"{category} ({len(cat_items)})", 2)
        p = doc.add_paragraph(CATEGORY_DESCRIPTIONS[category])
        p.italic = True
        
        for idx, item in enumerate(cat_items, 1):
            symbol = '✓' if category == 'WAHR' else '■'
            p = doc.add_paragraph()
            p.add_run(f'{idx}. {symbol} ').bold = True
            p.add_run(f'"{item["aussage"]}"').bold = True
            doc.add_paragraph(f"– {item['begründung']}")
    
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# ============================================
# SESSION STATE
# ============================================
if 'analysis_data' not in st.session_state:
    st.session_state.analysis_data = None
if 'summary' not in st.session_state:
    st.session_state.summary = None
if 'selected_model' not in st.session_state:
    st.session_state.selected_model = None
if 'input_text' not in st.session_state:
    st.session_state.input_text = ""

# ============================================
# SIDEBAR
# ============================================
with st.sidebar:
    st.title("⚙️ Einstellungen")
    
    # Load Models
    valid_models = fetch_models()
    
    if valid_models:
        st.subheader("AI Model")
        model_options = {f"{m['modelName']} ({m['provider']})": m for m in valid_models}
        selected_model_label = st.selectbox(
            "Wähle ein Model:",
            options=list(model_options.keys()),
            index=0
        )
        st.session_state.selected_model = model_options[selected_model_label]
        
        with st.expander("ℹ️ Model Details"):
            model = st.session_state.selected_model
            st.write(f"**ID:** {model['modelID']}")
            st.write(f"**Name:** {model['modelName']}")
            st.write(f"**Provider:** {model['provider']}")
    
    st.divider()
    
    st.subheader("📚 Kategorien")
    st.markdown("""
    - 🔴 **FALSCH** (5 Punkte)
    - 🟠 **DELEGITIMIERUNG** (4 Punkte)
    - 🟡 **VERZERRUNG** (3 Punkte)
    - 🟢 **FRAME** (1 Punkt)
    - 🔵 **WAHR** (0 Punkte)
    """)
    
    if st.session_state.analysis_data is not None:
        st.divider()
        if st.button("🔄 Neue Analyse", use_container_width=True):
            st.session_state.analysis_data = None
            st.session_state.summary = None
            st.session_state.input_text = ""
            st.rerun()

# ============================================
# MAIN APP
# ============================================

if not valid_models:
    st.error("⚠️ Keine Models verfügbar. Bitte API-Verbindung prüfen.")
    st.stop()

# Header
st.markdown('<div class="main-header">DESINFO</div>', unsafe_allow_html=True)
st.markdown('<div class="subtitle">Political Discourse Analyzer</div>', unsafe_allow_html=True)

# Main Content
if st.session_state.analysis_data is None:
    # INPUT MODE
    st.markdown('<div class="content-box">', unsafe_allow_html=True)
    
    st.markdown('<div class="section-title">Politische Aussagen eingeben</div>', unsafe_allow_html=True)
    st.markdown('<p style="color: #666; text-align: center; margin-bottom: 1.5rem;">Eine Aussage pro Zeile oder mit | getrennt</p>', unsafe_allow_html=True)
    
    input_text = st.text_area(
        "Text",
        height=350,
        placeholder="Beispiel:\nDiese Bundesregierung riskiert gerade den Dritten Weltkrieg.\nNehmen Sie Polen. Natürlich kann auch Polen für uns eine Gefahr sein.",
        key="text_input",
        value=st.session_state.input_text,
        label_visibility="collapsed"
    )
    
    st.session_state.input_text = input_text
    
    # Analyze Button
    st.markdown("<br>", unsafe_allow_html=True)
    
    if st.button("🚀 Analyse starten", type="primary", use_container_width=True, disabled=not bool(input_text and input_text.strip())):
        if input_text and input_text.strip():
            statements = parse_statements(input_text)
            model_id = st.session_state.selected_model['modelID']
            model_name = st.session_state.selected_model['modelName']
            
            with st.spinner(f"🔍 Analysiere {len(statements)} Aussagen mit {model_name}..."):
                result = call_api(statements, model_id)
                
                if result['success']:
                    st.session_state.analysis_data = result['data']
                    st.session_state.summary = calculate_summary(result['data'])
                    st.balloons()
                    st.rerun()
                else:
                    st.error(f"❌ Fehler: {result['error']}")
    
    # Preview
    if input_text:
        st.markdown("---")
        statements = parse_statements(input_text)
        st.info(f"📊 **{len(statements)} Aussagen** werden analysiert")
    
    st.markdown('</div>', unsafe_allow_html=True)

else:
    # RESULTS MODE
    analysis_data = st.session_state.analysis_data
    summary = st.session_state.summary
    model_info = st.session_state.selected_model
    
    # Score Display
    st.markdown(f"""
    <div class="score-box">
        <div class="score-value">{summary['desinfo_score']}</div>
        <div class="score-grade">Grade {summary['grade']}</div>
        <div class="score-label">{summary['grade_label']}</div>
    </div>
    """, unsafe_allow_html=True)
    
    # Quantifizierung
    st.markdown('<div class="content-box">', unsafe_allow_html=True)
    st.markdown('<div class="section-title">📊 Quantifizierung</div>', unsafe_allow_html=True)
    
    cols = st.columns(5)
    categories = ['FALSCH', 'DELEGITIMIERUNG', 'VERZERRUNG', 'FRAME', 'WAHR']
    
    for idx, cat in enumerate(categories):
        count = summary['category_counts'].get(cat, 0)
        percentage = (count / summary['total_statements'] * 100) if summary['total_statements'] > 0 else 0
        hex_color, _ = get_category_color(cat)
        
        with cols[idx]:
            st.markdown(f"""
            <div class="metric-card" style="background: {hex_color};">
                <div class="metric-value">{count}</div>
                <div class="metric-label">{cat}</div>
                <div class="metric-percent">{percentage:.0f}%</div>
            </div>
            """, unsafe_allow_html=True)
    
    st.markdown('</div>', unsafe_allow_html=True)
    
    # Detaillierte Ergebnisse - DIREKT SICHTBAR
    st.markdown('<div class="content-box">', unsafe_allow_html=True)
    st.markdown('<div class="section-title">📋 Detaillierte Ergebnisse</div>', unsafe_allow_html=True)
    
    for category in categories:
        cat_items = [item for item in analysis_data if item['kategorie'] == category]
        if not cat_items:
            continue
        
        # Category Header
        hex_color, _ = get_category_color(category)
        st.markdown(f'<div class="category-header" style="border-bottom-color: {hex_color};">{category} ({len(cat_items)} Aussagen)</div>', unsafe_allow_html=True)
        
        # Category Description
        st.markdown(f'<div class="category-description">{CATEGORY_DESCRIPTIONS[category]}</div>', unsafe_allow_html=True)
        
        # Statements - DIREKT ANGEZEIGT
        for idx, item in enumerate(cat_items, 1):
            st.markdown(f"""
            <div class="statement-card" style="border-left-color: {hex_color};">
                <div class="statement-title">{idx}. "{item['aussage']}"</div>
                <span class="category-badge badge-{category.lower()}">{category}</span>
                <span class="category-badge" style="background: #6c757d;">Punkte: {item['punkte']}</span>
                <div class="statement-text" style="margin-top: 1rem;"><strong>Begründung:</strong> {item['begründung']}</div>
            </div>
            """, unsafe_allow_html=True)
    
    st.markdown('</div>', unsafe_allow_html=True)
    
    # Downloads
    st.markdown('<div class="content-box">', unsafe_allow_html=True)
    st.markdown('<div class="section-title">📥 Report Download</div>', unsafe_allow_html=True)
    
    col1, col2 = st.columns(2)
    
    with col1:
        if st.button("📄 PDF Report generieren", use_container_width=True):
            with st.spinner("Generiere PDF..."):
                pdf_buffer = generate_pdf_report(analysis_data, summary, model_info)
                st.download_button(
                    label="⬇️ PDF herunterladen",
                    data=pdf_buffer,
                    file_name=f"DesInfo_Report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf",
                    mime="application/pdf",
                    use_container_width=True
                )
    
    with col2:
        if st.button("📝 DOCX Report generieren", use_container_width=True):
            with st.spinner("Generiere DOCX..."):
                docx_buffer = generate_docx_report(analysis_data, summary, model_info)
                st.download_button(
                    label="⬇️ DOCX herunterladen",
                    data=docx_buffer,
                    file_name=f"DesInfo_Report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True
                )
    
    st.markdown('</div>', unsafe_allow_html=True)

# Footer
st.markdown("---")
st.markdown("""
<div style="text-align: center; color: white; padding: 2rem 0;">
    <p style="font-size: 1.2rem; font-weight: 600;">Powered by <a href="https://www.democracy-intelligence.de" target="_blank" style="color: #ffd93d; text-decoration: none;">Democracy Intelligence</a></p>
    <p style="font-size: 0.9rem; opacity: 0.8;">DESINFO Political Discourse Analyzer v2.2</p>
</div>
""", unsafe_allow_html=True)
from reportlab.lib.units import cm
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak, Table, TableStyle
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY, TA_LEFT
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont

# DOCX Generation
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ============================================
# KONFIGURATION
# ============================================
try:
    API_BASE_URL = st.secrets["API_BASE_URL"]
except:
    API_BASE_URL = "http://217.154.156.197:8003"

MODELS_ENDPOINT = f"{API_BASE_URL}/desInfo/models"
ANALYZE_ENDPOINT = f"{API_BASE_URL}/desInfo/generateReport"

# Kategorie Scoring
CATEGORY_POINTS = {
    'FALSCH': 5,
    'DELEGITIMIERUNG': 4,
    'VERZERRUNG': 3,
    'FRAME': 1,
    'WAHR': 0
}

# Score-Grade Mapping
SCORE_GRADES = {
    'A': (0.0, 0.4, 'wahr', 'Die Aussagen entsprechen nachprüfbaren Fakten und sind wahrheitsgetreu.'),
    'B': (0.5, 1.4, 'geframed', 'Die Aussagen enthalten wahre Kernaussagen, die durch normative Framings eingefärbt werden.'),
    'C': (1.5, 2.4, 'verzerrend', 'Die Aussagen enthalten formal richtige Kernaussagen, die durch Übertreibung, Verkürzung oder Kontextausblendung ein schiefes Bild erzeugen.'),
    'D': (2.5, 3.4, 'demokratisch dysfunktional', 'Die Aussagen enthalten substanzielle Falschbehauptungen und delegitimierende Elemente, die den demokratischen Diskurs belasten.'),
    'E': (3.5, 5.0, 'demokratisch destruktiv', 'Die Aussagen sind überwiegend falsch und delegitimierend, untergraben systematisch Vertrauen und demokratische Institutionen.')
}

CATEGORY_DESCRIPTIONS = {
    'FALSCH': 'Objektiv widerlegte Behauptungen. Es liegen belastbare Daten oder Ereignisprotokolle vor, die das Gegenteil zeigen. Keine Meinungen oder Prognosen; nur Tatsachenbehauptungen, die nachweislich nicht stimmen.',
    'DELEGITIMIERUNG': 'Abwertung oder Untergrabung von Personen, Gruppen oder Institutionen. Sprachliche Diskreditierungen, Kampfbegriffe oder systematische Diffamierungen.',
    'VERZERRUNG': 'Formal richtige Kernaussagen, die durch Übertreibung, Verkürzung oder Kontextausblendung ein schiefes Bild erzeugen.',
    'FRAME': 'Sprachliche Deutungsrahmen, die neutrale Sachverhalte emotional oder normativ aufladen, ohne sie faktisch falsch oder delegitimierend darzustellen.',
    'WAHR': 'Sachlich zutreffende Aussagen, die überprüfbar sind und keinen überzogenen Frame, keine Verzerrung oder Delegitimierung enthalten.'
}

# ============================================
# PAGE CONFIG
# ============================================
st.set_page_config(
    page_title="DESINFO Analyzer",
    page_icon="📊",
    layout="wide",
    initial_sidebar_state="collapsed"
)

# ============================================
# CUSTOM CSS - Democracy Intelligence Design
# ============================================
st.markdown("""
<style>
    @import url('https://fonts.googleapis.com/css2?family=Inter:wght@400;600;700;800&display=swap');
    
    /* Global Styles */
    * {
        font-family: 'Inter', -apple-system, BlinkMacSystemFont, sans-serif;
    }
    
    /* Main Container */
    .main {
        background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
        padding: 0;
    }
    
    .block-container {
        max-width: 1200px;
        padding: 2rem 1rem;
    }
    
    /* Header */
    .main-header {
        font-size: 3rem;
        font-weight: 800;
        color: white;
        text-align: center;
        padding: 3rem 0 1rem 0;
        margin-bottom: 0.5rem;
    }
    
    .subtitle {
        font-size: 1.4rem;
        color: #ffd93d;
        text-align: center;
        font-weight: 600;
        margin-bottom: 3rem;
    }
    
    /* Content Box */
    .content-box {
        background: white;
        border-radius: 20px;
        padding: 2rem;
        margin: 2rem 0;
        box-shadow: 0 10px 40px rgba(0,0,0,0.2);
    }
    
    /* Buttons */
    .stButton>button {
        background: linear-gradient(135deg, #ffd93d 0%, #ffb800 100%);
        color: #1a1a1a;
        border: none;
        padding: 0.75rem 2rem;
        font-size: 1.1rem;
        font-weight: 700;
        border-radius: 50px;
        transition: all 0.3s;
        width: 100%;
    }
    
    .stButton>button:hover {
        transform: translateY(-2px);
        box-shadow: 0 6px 20px rgba(255, 217, 61, 0.4);
        background: linear-gradient(135deg, #ffb800 0%, #ffd93d 100%);
    }
    
    /* Text Area */
    .stTextArea textarea {
        border: 2px solid #667eea;
        border-radius: 15px;
        font-size: 1rem;
        padding: 1rem;
    }
    
    /* Category Badges */
    .category-badge {
        display: inline-block;
        padding: 0.5rem 1.2rem;
        border-radius: 25px;
        font-weight: 700;
        font-size: 0.9rem;
        margin: 0.3rem;
    }
    
    .badge-falsch { background: #dc3545; color: white; }
    .badge-delegitimierung { background: #fd7e14; color: white; }
    .badge-verzerrung { background: #ffc107; color: #000; }
    .badge-frame { background: #28a745; color: white; }
    .badge-wahr { background: #007bff; color: white; }
    
    /* Score Display */
    .score-box {
        background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
        border-radius: 20px;
        padding: 3rem 2rem;
        text-align: center;
        color: white;
        margin: 2rem 0;
        box-shadow: 0 10px 30px rgba(102, 126, 234, 0.3);
    }
    
    .score-value {
        font-size: 6rem;
        font-weight: 800;
        margin: 0;
        text-shadow: 0 4px 10px rgba(0,0,0,0.2);
    }
    
    .score-grade {
        font-size: 2.5rem;
        font-weight: 700;
        margin: 1rem 0;
        color: #ffd93d;
    }
    
    .score-label {
        font-size: 1.8rem;
        opacity: 0.95;
        font-weight: 600;
    }
    
    /* Metric Cards */
    .metric-card {
        text-align: center;
        padding: 2rem 1rem;
        border-radius: 15px;
        color: white;
        margin: 0.5rem;
        box-shadow: 0 5px 15px rgba(0,0,0,0.2);
        transition: transform 0.3s;
    }
    
    .metric-card:hover {
        transform: translateY(-5px);
    }
    
    .metric-value {
        font-size: 3rem;
        font-weight: 800;
        margin: 0.5rem 0;
    }
    
    .metric-label {
        font-size: 1.1rem;
        font-weight: 600;
        opacity: 0.9;
    }
    
    .metric-percent {
        font-size: 1rem;
        opacity: 0.8;
    }
    
    /* Statement Card */
    .statement-card {
        background: #f8f9fa;
        border-left: 5px solid #667eea;
        padding: 1.5rem;
        margin: 1.5rem 0;
        border-radius: 10px;
        box-shadow: 0 2px 8px rgba(0,0,0,0.1);
    }
    
    /* Select Box */
    .stSelectbox {
        background: white;
        border-radius: 10px;
    }
    
    /* Hide Streamlit Branding */
    #MainMenu {visibility: hidden;}
    footer {visibility: hidden;}
    
    /* Section Headers */
    .section-header {
        font-size: 2rem;
        font-weight: 700;
        color: white;
        text-align: center;
        margin: 2rem 0 1rem 0;
    }
    
    .section-header-dark {
        font-size: 2rem;
        font-weight: 700;
        color: #1a1a1a;
        text-align: center;
        margin: 2rem 0 1rem 0;
    }
</style>
""", unsafe_allow_html=True)

# ============================================
# HELPER FUNCTIONS
# ============================================

@st.cache_data(ttl=300)
def fetch_models():
    try:
        response = requests.get(MODELS_ENDPOINT, timeout=10)
        if response.ok:
            data = response.json()
            return [m for m in data.get('models', []) if m.get('valid', False)]
    except Exception as e:
        st.error(f"⚠️ Fehler beim Laden der Models: {e}")
    return []

def parse_statements(text: str) -> list:
    if '|' in text:
        statements = [s.strip() for s in text.split('|') if s.strip()]
    else:
        statements = [s.strip() for s in text.split('\n') if s.strip()]
    return statements

def call_api(statements: list, model_id: int) -> dict:
    text_param = "|".join(statements)
    params = {"modelID": model_id, "text": text_param}
    
    try:
        response = requests.get(ANALYZE_ENDPOINT, params=params, timeout=120)
        response.raise_for_status()
        return {"success": True, "data": response.json()}
    except Exception as e:
        return {"success": False, "error": str(e)}

def calculate_summary(analysis_data: list) -> dict:
    category_counts = Counter(item['kategorie'] for item in analysis_data)
    total_points = sum(CATEGORY_POINTS.get(item['kategorie'], 0) for item in analysis_data)
    total_statements = len(analysis_data)
    desinfo_score = total_points / total_statements if total_statements > 0 else 0
    
    grade = 'E'
    grade_label = ''
    grade_description = ''
    
    for g, (min_score, max_score, label, desc) in SCORE_GRADES.items():
        if min_score <= desinfo_score <= max_score:
            grade = g
            grade_label = label
            grade_description = desc
            break
    
    return {
        'total_statements': total_statements,
        'category_counts': dict(category_counts),
        'total_points': total_points,
        'desinfo_score': round(desinfo_score, 1),
        'grade': grade,
        'grade_label': grade_label,
        'grade_description': grade_description
    }

def get_category_color(category: str) -> tuple:
    """Returns (hex_color, rgb_color) for ReportLab"""
    color_map = {
        'FALSCH': ('#dc3545', colors.HexColor('#dc3545')),
        'DELEGITIMIERUNG': ('#fd7e14', colors.HexColor('#fd7e14')),
        'VERZERRUNG': ('#ffc107', colors.HexColor('#ffc107')),
        'FRAME': ('#28a745', colors.HexColor('#28a745')),
        'WAHR': ('#007bff', colors.HexColor('#007bff'))
    }
    return color_map.get(category, ('#6c757d', colors.grey))

def generate_pdf_report(analysis_data: list, summary: dict, model_info: dict = None) -> BytesIO:
    """Generiert PDF Report - EXAKT wie Beispiel"""
    buffer = BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=A4,
        topMargin=2*cm,
        bottomMargin=2*cm,
        leftMargin=2.5*cm,
        rightMargin=2.5*cm
    )
    
    elements = []
    styles = getSampleStyleSheet()
    
    # Custom Styles
    title_style = ParagraphStyle(
        'Title',
        parent=styles['Heading1'],
        fontSize=24,
        leading=28,
        textColor=colors.black,
        spaceAfter=20,
        alignment=TA_LEFT,
        fontName='Helvetica-Bold'
    )
    
    heading_style = ParagraphStyle(
        'Heading',
        parent=styles['Heading2'],
        fontSize=16,
        leading=20,
        textColor=colors.black,
        spaceAfter=10,
        spaceBefore=15,
        fontName='Helvetica-Bold'
    )
    
    body_style = ParagraphStyle(
        'Body',
        parent=styles['Normal'],
        fontSize=11,
        leading=15,
        alignment=TA_JUSTIFY,
        fontName='Helvetica'
    )
    
    italic_style = ParagraphStyle(
        'Italic',
        parent=body_style,
        fontName='Helvetica-Oblique'
    )
    
    # ====== SEITE 1 ======
    
    # Titel
    elements.append(Paragraph("Vollständige Auswertung", title_style))
    elements.append(Spacer(1, 0.3*cm))
    
    # Linie
    line = Table([['']], colWidths=[16*cm])
    line.setStyle(TableStyle([('LINEABOVE', (0,0), (-1,0), 1, colors.grey)]))
    elements.append(line)
    elements.append(Spacer(1, 0.8*cm))
    
    # Zusammenfassung
    elements.append(Paragraph("Zusammenfassung", heading_style))
    
    total = summary['total_statements']
    counts = summary['category_counts']
    
    summary_text = f"Die vorliegende Analyse untersucht {total} Aussagen nach ihrer faktischen Richtigkeit und kommunikativen Qualität. Das Ergebnis zeigt eine ausgeprägte Tendenz zu {summary['grade_label']}en Darstellungen und normativen Framings, während objektiv falsche Behauptungen und delegitimierende Aussagen in moderatem Umfang vorkommen. Wahrheitsgetreue Aussagen machen nur einen kleinen Teil aus."
    
    elements.append(Paragraph(summary_text, body_style))
    elements.append(Spacer(1, 0.8*cm))
    elements.append(line)
    elements.append(Spacer(1, 0.8*cm))
    
    # Quantifizierung
    elements.append(Paragraph("Quantifizierung", heading_style))
    elements.append(Paragraph(f"Anzahl Aussagen gesamt: {total}", body_style))
    elements.append(Spacer(1, 0.3*cm))
    
    # Kategorien mit Farben
    for cat in ['FALSCH', 'DELEGITIMIERUNG', 'VERZERRUNG', 'FRAME', 'WAHR']:
        count = counts.get(cat, 0)
        percentage = (count / total * 100) if total > 0 else 0
        hex_color, _ = get_category_color(cat)
        
        # Colored square bullet
        cat_text = f'<font color="{hex_color}">■</font> {cat}: {count} ({percentage:.0f} %)'
        elements.append(Paragraph(cat_text, body_style))
    
    elements.append(Spacer(1, 0.8*cm))
    elements.append(line)
    elements.append(Spacer(1, 0.8*cm))
    
    # Scoring
    elements.append(Paragraph("Scoring", heading_style))
    score_text = f"<b>Desinfo-Score: {summary['desinfo_score']}</b><br/>{summary['grade']}: {summary['grade_label']}<br/><br/>{summary['grade_description']}"
    elements.append(Paragraph(score_text, body_style))
    elements.append(Spacer(1, 0.8*cm))
    elements.append(line)
    elements.append(Spacer(1, 0.8*cm))
    
    # Einteilung
    elements.append(Paragraph("Einteilung", heading_style))
    for grade, (min_s, max_s, label, _) in SCORE_GRADES.items():
        elements.append(Paragraph(f"{grade}: {min_s} bis {max_s}: {label}", body_style))
    
    # Page Break
    elements.append(PageBreak())
    
    # ====== SEITE 2+ ======
    
    elements.append(Paragraph("Vollständige Auswertung des Textes", title_style))
    elements.append(Spacer(1, 0.3*cm))
    elements.append(line)
    elements.append(Spacer(1, 0.8*cm))
    
    # Kategorien
    categories_order = ['FALSCH', 'DELEGITIMIERUNG', 'VERZERRUNG', 'FRAME', 'WAHR']
    
    for category in categories_order:
        cat_items = [item for item in analysis_data if item['kategorie'] == category]
        if not cat_items:
            continue
        
        # Category Header
        elements.append(Paragraph(f"<b>{category} ({len(cat_items)})</b>", heading_style))
        
        # Description
        elements.append(Paragraph(CATEGORY_DESCRIPTIONS[category], italic_style))
        elements.append(Spacer(1, 0.5*cm))
        
        # Statements
        for idx, item in enumerate(cat_items, 1):
            hex_color, _ = get_category_color(category)
            symbol = '✓' if category == 'WAHR' else '■'
            
            # Statement
            stmt_text = f'{idx}. <font color="{hex_color}">{symbol}</font> <b>"{item["aussage"]}"</b>'
            elements.append(Paragraph(stmt_text, body_style))
            elements.append(Spacer(1, 0.2*cm))
            
            # Begründung
            beg_text = f"– {item['begründung']}"
            elements.append(Paragraph(beg_text, body_style))
            elements.append(Spacer(1, 0.5*cm))
        
        elements.append(Spacer(1, 0.5*cm))
    
    # Build
    doc.build(elements)
    buffer.seek(0)
    return buffer

def generate_docx_report(analysis_data: list, summary: dict, model_info: dict = None) -> BytesIO:
    """Generiert DOCX Report - EXAKT wie PDF"""
    doc = Document()
    
    # Titel
    title = doc.add_heading('Vollständige Auswertung', 0)
    
    # Zusammenfassung
    doc.add_heading('Zusammenfassung', 1)
    total = summary['total_statements']
    summary_text = f"Die vorliegende Analyse untersucht {total} Aussagen nach ihrer faktischen Richtigkeit und kommunikativen Qualität. Das Ergebnis zeigt eine ausgeprägte Tendenz zu {summary['grade_label']}en Darstellungen."
    doc.add_paragraph(summary_text)
    
    # Quantifizierung
    doc.add_heading('Quantifizierung', 1)
    doc.add_paragraph(f"Anzahl Aussagen gesamt: {total}")
    
    counts = summary['category_counts']
    for cat in ['FALSCH', 'DELEGITIMIERUNG', 'VERZERRUNG', 'FRAME', 'WAHR']:
        count = counts.get(cat, 0)
        percentage = (count / total * 100) if total > 0 else 0
        doc.add_paragraph(f"■ {cat}: {count} ({percentage:.0f} %)")
    
    # Scoring
    doc.add_heading('Scoring', 1)
    doc.add_paragraph(f"Desinfo-Score: {summary['desinfo_score']}")
    doc.add_paragraph(f"{summary['grade']}: {summary['grade_label']}")
    doc.add_paragraph(summary['grade_description'])
    
    # Einteilung
    doc.add_heading('Einteilung', 1)
    for grade, (min_s, max_s, label, _) in SCORE_GRADES.items():
        doc.add_paragraph(f"{grade}: {min_s} bis {max_s}: {label}")
    
    # Page Break
    doc.add_page_break()
    doc.add_heading('Vollständige Auswertung des Textes', 1)
    
    # Categories
    categories_order = ['FALSCH', 'DELEGITIMIERUNG', 'VERZERRUNG', 'FRAME', 'WAHR']
    
    for category in categories_order:
        cat_items = [item for item in analysis_data if item['kategorie'] == category]
        if not cat_items:
            continue
        
        doc.add_heading(f"{category} ({len(cat_items)})", 2)
        
        p = doc.add_paragraph(CATEGORY_DESCRIPTIONS[category])
        p.italic = True
        
        for idx, item in enumerate(cat_items, 1):
            symbol = '✓' if category == 'WAHR' else '■'
            p = doc.add_paragraph()
            p.add_run(f'{idx}. {symbol} ').bold = True
            p.add_run(f'"{item["aussage"]}"').bold = True
            
            doc.add_paragraph(f"– {item['begründung']}")
    
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# ============================================
# SESSION STATE
# ============================================
if 'analysis_data' not in st.session_state:
    st.session_state.analysis_data = None
if 'summary' not in st.session_state:
    st.session_state.summary = None
if 'selected_model' not in st.session_state:
    st.session_state.selected_model = None
if 'input_text' not in st.session_state:
    st.session_state.input_text = ""

# ============================================
# LOAD MODELS
# ============================================
valid_models = fetch_models()

if not valid_models:
    st.error("⚠️ Keine Models verfügbar. Bitte API-Verbindung prüfen.")
    st.stop()

# ============================================
# MAIN APP
# ============================================

# Header
st.markdown('<div class="main-header">DESINFO</div>', unsafe_allow_html=True)
st.markdown('<div class="subtitle">Political Discourse Analyzer</div>', unsafe_allow_html=True)

# Main Content
if st.session_state.analysis_data is None:
    # INPUT MODE
    st.markdown('<div class="content-box">', unsafe_allow_html=True)
    
    # Model Selection
    st.markdown('<h3 style="color: #1a1a1a; text-align: center; margin-bottom: 1rem;">Wähle dein AI Model</h3>', unsafe_allow_html=True)
    
    model_options = {f"{m['modelName']} ({m['provider']})": m for m in valid_models}
    selected_model_label = st.selectbox(
        "Model",
        options=list(model_options.keys()),
        index=0,
        label_visibility="collapsed"
    )
    st.session_state.selected_model = model_options[selected_model_label]
    
    st.markdown("---")
    
    # Text Input
    st.markdown('<h3 style="color: #1a1a1a; margin-top: 2rem;">Politische Aussagen eingeben</h3>', unsafe_allow_html=True)
    st.markdown('<p style="color: #666; margin-bottom: 1rem;">Eine Aussage pro Zeile oder mit | getrennt</p>', unsafe_allow_html=True)
    
    input_text = st.text_area(
        "Text",
        height=300,
        placeholder="Beispiel:\nDiese Bundesregierung riskiert gerade den Dritten Weltkrieg.\nNehmen Sie Polen. Natürlich kann auch Polen für uns eine Gefahr sein.",
        key="text_input",
        value=st.session_state.input_text,
        label_visibility="collapsed"
    )
    
    st.session_state.input_text = input_text
    
    # Buttons Row
    col1, col2 = st.columns(2)
    
    with col1:
        if st.button("📝 Beispiel laden", use_container_width=True):
            st.session_state.input_text = """Diese Bundesregierung riskiert gerade den Dritten Weltkrieg.|Nehmen Sie Polen. Natürlich kann auch Polen für uns eine Gefahr sein.|Das macht Polen nicht."""
            st.rerun()
    
    with col2:
        # FIX: Button IMMER enabled wenn Text da ist
        if st.button("🚀 Analyse starten", type="primary", use_container_width=True, disabled=not bool(input_text and input_text.strip())):
            if input_text and input_text.strip():
                statements = parse_statements(input_text)
                model_id = st.session_state.selected_model['modelID']
                model_name = st.session_state.selected_model['modelName']
                
                with st.spinner(f"🔍 Analysiere mit {model_name}..."):
                    result = call_api(statements, model_id)
                    
                    if result['success']:
                        st.session_state.analysis_data = result['data']
                        st.session_state.summary = calculate_summary(result['data'])
                        st.balloons()
                        st.rerun()
                    else:
                        st.error(f"❌ Fehler: {result['error']}")
    
    # Preview
    if input_text:
        with st.expander("👁️ Statement Preview"):
            statements = parse_statements(input_text)
            st.metric("Anzahl Statements", len(statements))
            for i, stmt in enumerate(statements, 1):
                st.write(f"{i}. {stmt}")
    
    st.markdown('</div>', unsafe_allow_html=True)

else:
    # RESULTS MODE
    analysis_data = st.session_state.analysis_data
    summary = st.session_state.summary
    model_info = st.session_state.selected_model
    
    # Score Display
    st.markdown(f"""
    <div class="score-box">
        <div class="score-value">{summary['desinfo_score']}</div>
        <div class="score-grade">Grade {summary['grade']}</div>
        <div class="score-label">{summary['grade_label']}</div>
    </div>
    """, unsafe_allow_html=True)
    
    # Quantifizierung
    st.markdown('<div class="content-box">', unsafe_allow_html=True)
    st.markdown('<h2 class="section-header-dark">📊 Quantifizierung</h2>', unsafe_allow_html=True)
    
    cols = st.columns(5)
    categories = ['FALSCH', 'DELEGITIMIERUNG', 'VERZERRUNG', 'FRAME', 'WAHR']
    
    for idx, cat in enumerate(categories):
        count = summary['category_counts'].get(cat, 0)
        percentage = (count / summary['total_statements'] * 100) if summary['total_statements'] > 0 else 0
        hex_color, _ = get_category_color(cat)
        
        with cols[idx]:
            st.markdown(f"""
            <div class="metric-card" style="background: {hex_color};">
                <div class="metric-value">{count}</div>
                <div class="metric-label">{cat}</div>
                <div class="metric-percent">{percentage:.0f}%</div>
            </div>
            """, unsafe_allow_html=True)
    
    st.markdown('</div>', unsafe_allow_html=True)
    
    # Detailed Results
    st.markdown('<div class="content-box">', unsafe_allow_html=True)
    st.markdown('<h2 class="section-header-dark">📋 Detaillierte Ergebnisse</h2>', unsafe_allow_html=True)
    
    for category in categories:
        cat_items = [item for item in analysis_data if item['kategorie'] == category]
        if not cat_items:
            continue
        
        with st.expander(f"**{category}** ({len(cat_items)} Aussagen)", expanded=False):
            st.markdown(f"*{CATEGORY_DESCRIPTIONS[category]}*")
            st.markdown("---")
            
            for idx, item in enumerate(cat_items, 1):
                hex_color, _ = get_category_color(category)
                st.markdown(f"""
                <div class="statement-card" style="border-left-color: {hex_color};">
                    <h4>{idx}. "{item['aussage']}"</h4>
                    <span class="category-badge badge-{category.lower()}">{category}</span>
                    <span class="category-badge" style="background: #6c757d; color: white;">Punkte: {item['punkte']}</span>
                    <p style="margin-top: 1rem;"><strong>Begründung:</strong> {item['begründung']}</p>
                </div>
                """, unsafe_allow_html=True)
    
    st.markdown('</div>', unsafe_allow_html=True)
    
    # Downloads
    st.markdown('<div class="content-box">', unsafe_allow_html=True)
    st.markdown('<h2 class="section-header-dark">📥 Report Download</h2>', unsafe_allow_html=True)
    
    col1, col2, col3 = st.columns([1, 1, 1])
    
    with col1:
        if st.button("📄 PDF Report generieren", use_container_width=True):
            with st.spinner("Generiere PDF..."):
                pdf_buffer = generate_pdf_report(analysis_data, summary, model_info)
                st.download_button(
                    label="⬇️ PDF herunterladen",
                    data=pdf_buffer,
                    file_name=f"DesInfo_Report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf",
                    mime="application/pdf",
                    use_container_width=True
                )
    
    with col2:
        if st.button("📝 DOCX Report generieren", use_container_width=True):
            with st.spinner("Generiere DOCX..."):
                docx_buffer = generate_docx_report(analysis_data, summary, model_info)
                st.download_button(
                    label="⬇️ DOCX herunterladen",
                    data=docx_buffer,
                    file_name=f"DesInfo_Report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True
                )
    
    with col3:
        if st.button("🔄 Neue Analyse", use_container_width=True):
            st.session_state.analysis_data = None
            st.session_state.summary = None
            st.session_state.input_text = ""
            st.rerun()
    
    st.markdown('</div>', unsafe_allow_html=True)

# Footer
st.markdown("---")
st.markdown("""
<div style="text-align: center; color: white; padding: 2rem 0;">
    <p style="font-size: 1.2rem; font-weight: 600;">Powered by <a href="https://www.democracy-intelligence.de" target="_blank" style="color: #ffd93d; text-decoration: none;">Democracy Intelligence</a></p>
    <p style="font-size: 0.9rem; opacity: 0.8;">DESINFO Political Discourse Analyzer v2.1</p>
</div>
""", unsafe_allow_html=True)
